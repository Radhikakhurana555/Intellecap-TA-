import os
from io import BytesIO
from typing import Optional, Dict, List

import pandas as pd
import streamlit as st
from openai import OpenAI
from pypdf import PdfReader
from docx import Document


st.set_page_config(
    page_title="Intellecap TA Copilot",
    page_icon="🤖",
    layout="wide"
)

APP_TITLE = "Intellecap TA Copilot"
APP_SUBTITLE = "Resume screening, candidate summaries, interview questions, and stakeholder-ready hiring updates."


def get_api_key() -> Optional[str]:
    """Read API key from Streamlit secrets or environment variable."""
    try:
        return st.secrets.get("OPENAI_API_KEY")
    except Exception:
        return os.getenv("OPENAI_API_KEY")


def extract_text_from_pdf(file) -> str:
    reader = PdfReader(file)
    pages = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages.append(text)
    return "\n".join(pages).strip()


def extract_text_from_docx(file) -> str:
    document = Document(file)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    tables = []
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                tables.append(row_text)
    return "\n".join(paragraphs + tables).strip()


def extract_text(file) -> str:
    """Extract text from PDF, DOCX, or TXT files."""
    file_name = file.name.lower()
    if file_name.endswith(".pdf"):
        return extract_text_from_pdf(file)
    if file_name.endswith(".docx"):
        return extract_text_from_docx(file)
    if file_name.endswith(".txt"):
        return file.getvalue().decode("utf-8", errors="ignore").strip()
    return ""


def truncate_text(text: str, limit: int = 20000) -> str:
    """Keep prompt size controlled."""
    if len(text) <= limit:
        return text
    return text[:limit] + "\n\n[Text truncated for analysis. Please review original document if needed.]"


def build_prompt(
    jd_text: str,
    resume_text: str,
    role_title: str,
    candidate_context: str,
    screening_weights: str,
) -> str:
    return f"""
You are an HR Talent Acquisition Copilot for Intellecap Advisory Services, an impact consulting and advisory organization.

Your task is to assess the candidate against the job description using only the information provided below.
Do not infer age, gender, caste, religion, marital status, disability, nationality, or any other protected/personal attribute.
Do not make a final hiring decision. Provide a structured recommendation for human review.

Role title:
{role_title}

Screening priorities / weights:
{screening_weights}

Additional candidate context, if provided:
{candidate_context}

JOB DESCRIPTION:
{truncate_text(jd_text, 16000)}

CANDIDATE RESUME:
{truncate_text(resume_text, 22000)}

Return the response in the following format:

## 1. Fitment Snapshot
- Fitment Score: __/100
- Recommendation: Shortlist / Hold / Not aligned
- Confidence Level: High / Medium / Low
- One-line Rationale:

## 2. Crisp Candidate Summary
Write a stakeholder-ready summary in 80-100 words. Keep the tone neutral and professional.

## 3. Role Fitment
Create a table with:
- Requirement
- Evidence from Resume
- Fitment: Strong / Moderate / Weak / Not Evident

## 4. Strengths
Provide 4-6 crisp bullets.

## 5. Possible Concerns / Gaps
Provide 3-5 crisp bullets. Use neutral wording.

## 6. Suggested Interview Questions
Provide 6 questions:
- 3 technical/domain questions
- 2 behavioral/stakeholder management questions
- 1 compensation/motivation or availability question

## 7. Stakeholder Email Draft
Write a concise email that HR can send to the hiring manager with the candidate summary, fitment, and suggested next step.

## 8. Human Review Note
Mention that this is AI-assisted screening and should be validated by HR and the hiring manager.
"""


def call_openai(prompt: str, model_name: str, api_key: str) -> str:
    client = OpenAI(api_key=api_key)
    response = client.responses.create(
        model=model_name,
        instructions=(
            "You are a careful HR assistant. Be concise, neutral, and evidence-based. "
            "Avoid discriminatory or protected-class inferences. Use only the supplied text."
        ),
        input=prompt,
    )
    return response.output_text


def make_download_bundle(results: Dict[str, str]) -> bytes:
    output = BytesIO()
    with zipfile.ZipFile(output, "w", zipfile.ZIP_DEFLATED) as zf:
        for file_name, content in results.items():
            safe_name = "".join(ch if ch.isalnum() or ch in (" ", "-", "_", ".") else "_" for ch in file_name)
            zf.writestr(f"{safe_name}.md", content)
    output.seek(0)
    return output.getvalue()


def parse_score(result_text: str) -> Optional[int]:
    import re
    patterns = [
        r"Fitment Score:\s*(\d{1,3})\s*/\s*100",
        r"Fitment Score\s*[:\-]\s*(\d{1,3})",
        r"Score\s*[:\-]\s*(\d{1,3})\s*/\s*100",
    ]
    for pattern in patterns:
        match = re.search(pattern, result_text, flags=re.IGNORECASE)
        if match:
            score = int(match.group(1))
            return max(0, min(score, 100))
    return None


def main():
    st.title(APP_TITLE)
    st.caption(APP_SUBTITLE)

    with st.sidebar:
        st.header("Settings")
        role_title = st.text_input("Role title", value="Consultant / Manager")
        model_name = st.text_input("OpenAI model", value="gpt-5.2")
        st.info(
            "Keep the app human-in-the-loop. The bot assists screening and drafting; HR and hiring managers should make final decisions."
        )

        st.subheader("Screening priorities")
        domain_weight = st.slider("Domain relevance", 0, 100, 30)
        consulting_weight = st.slider("Consulting / advisory experience", 0, 100, 25)
        stakeholder_weight = st.slider("Stakeholder management", 0, 100, 20)
        communication_weight = st.slider("Communication / writing", 0, 100, 15)
        location_weight = st.slider("Location / availability / compensation", 0, 100, 10)

        screening_weights = f"""
        Domain relevance: {domain_weight}%
        Consulting/advisory experience: {consulting_weight}%
        Stakeholder management: {stakeholder_weight}%
        Communication/writing: {communication_weight}%
        Location/availability/compensation: {location_weight}%
        """

    api_key = get_api_key()

    if not api_key:
        st.warning(
            "OPENAI_API_KEY is not configured. Add it in Streamlit Cloud → App settings → Secrets. "
            "You can still prepare the inputs, but analysis will run after the key is added."
        )

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("1. Upload Job Description")
        jd_file = st.file_uploader("Upload JD file", type=["pdf", "docx", "txt"], key="jd")
        jd_manual = st.text_area("Or paste JD here", height=240)

    with col2:
        st.subheader("2. Upload Candidate Resumes")
        resume_files = st.file_uploader(
            "Upload one or multiple resumes",
            type=["pdf", "docx", "txt"],
            accept_multiple_files=True,
            key="resumes",
        )
        candidate_context = st.text_area(
            "Optional candidate context",
            placeholder="Example: Current CTC, expected CTC, notice period, location preference, interview notes...",
            height=240,
        )

    jd_text = ""
    if jd_file:
        jd_text = extract_text(jd_file)
    elif jd_manual.strip():
        jd_text = jd_manual.strip()

    with st.expander("Preview extracted JD text"):
        st.write(jd_text[:4000] if jd_text else "No JD text available yet.")

    st.divider()

    run_button = st.button("Run TA Screening", type="primary", use_container_width=True)

    if run_button:
        if not jd_text:
            st.error("Please upload or paste the job description.")
            return
        if not resume_files:
            st.error("Please upload at least one resume.")
            return
        if not api_key:
            st.error("Please configure OPENAI_API_KEY in Streamlit secrets before running analysis.")
            return

        all_results: Dict[str, str] = {}
        summary_rows: List[Dict[str, str]] = []

        progress = st.progress(0)
        status = st.empty()

        for idx, resume_file in enumerate(resume_files, start=1):
            status.write(f"Analyzing {resume_file.name}...")
            resume_text = extract_text(resume_file)

            if not resume_text:
                result = f"Could not extract readable text from {resume_file.name}. Please upload a text-readable PDF/DOCX/TXT file."
            else:
                prompt = build_prompt(
                    jd_text=jd_text,
                    resume_text=resume_text,
                    role_title=role_title,
                    candidate_context=candidate_context,
                    screening_weights=screening_weights,
                )
                try:
                    result = call_openai(prompt, model_name=model_name, api_key=api_key)
                except Exception as exc:
                    result = f"Analysis failed for {resume_file.name}: {exc}"

            all_results[resume_file.name] = result
            score = parse_score(result)

            recommendation = ""
            for line in result.splitlines():
                if "Recommendation:" in line:
                    recommendation = line.split("Recommendation:", 1)[-1].strip()
                    break

            summary_rows.append({
                "Candidate File": resume_file.name,
                "Fitment Score": score if score is not None else "",
                "Recommendation": recommendation,
            })

            with st.expander(f"Result: {resume_file.name}", expanded=True):
                st.markdown(result)
                st.download_button(
                    label=f"Download analysis for {resume_file.name}",
                    data=result,
                    file_name=f"{resume_file.name}_analysis.md",
                    mime="text/markdown",
                )

            progress.progress(idx / len(resume_files))

        status.success("Analysis completed.")

        st.subheader("Candidate Comparison Snapshot")
        summary_df = pd.DataFrame(summary_rows)
        st.dataframe(summary_df, use_container_width=True)

        csv_data = summary_df.to_csv(index=False).encode("utf-8")
        st.download_button(
            "Download comparison CSV",
            data=csv_data,
            file_name="candidate_comparison.csv",
            mime="text/csv",
        )

        st.download_button(
            "Download all analyses as ZIP",
            data=make_download_bundle(all_results),
            file_name="ta_copilot_candidate_analyses.zip",
            mime="application/zip",
        )


if __name__ == "__main__":
    main()
